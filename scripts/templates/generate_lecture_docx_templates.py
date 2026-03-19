#!/usr/bin/env python3
"""Generate parser-safe DOCX templates for lecture-organized CAH imports."""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document


def lecture_code(index: int) -> str:
    return f"{index:02d}"


def add_intro(doc: Document, lecture: str) -> None:
    doc.add_heading(f"CAH QBank MCQ Template - Lecture {lecture}", level=1)
    doc.add_paragraph("Use this template to add paediatric MCQs or EMQ stems for one teaching block.")
    doc.add_paragraph("Allowed types: SBA and EMQ stem only.")
    doc.add_paragraph("Keep question numbering unique inside this file and complete the final Answer Key block.")
    doc.add_paragraph("Tags should include CAH module/topic labels when known, e.g. CAH 05 > Respiratory; Lecture > 12")


def add_emq_block(doc: Document, lecture: str) -> None:
    doc.add_paragraph("Section A - Extended Matching Questions (EMQ)")
    doc.add_paragraph("EMQ Set A: [Replace with EMQ set title]")
    doc.add_paragraph("Select the SINGLE best option for each question.")
    doc.add_paragraph("Options:")
    doc.add_paragraph("A. [EMQ option A]")
    doc.add_paragraph("B. [EMQ option B]")
    doc.add_paragraph("C. [EMQ option C]")
    doc.add_paragraph("D. [EMQ option D]")
    doc.add_paragraph("E. [EMQ option E]")
    doc.add_paragraph("Question 1. [EMQ stem 1]")
    doc.add_paragraph(
        f"Tags: CAH 01 > [Subtopic]; Lecture > {lecture}; Difficulty: Intermediate; AUS:3; Domain: [Curriculum area]; Yield: [High/Medium/Low]"
    )
    doc.add_paragraph("Question 2. [EMQ stem 2]")
    doc.add_paragraph(
        f"Tags: CAH 01 > [Subtopic]; Lecture > {lecture}; Difficulty: Intermediate; AUS:3; Domain: [Curriculum area]; Yield: [High/Medium/Low]"
    )


def add_sba_block(doc: Document, lecture: str) -> None:
    doc.add_paragraph("Section B - Single Best Answer (SBA)")
    for q in range(3, 11):
        doc.add_paragraph(f"Question {q}. [SBA stem {q}]")
        doc.add_paragraph("A. [Option A]")
        doc.add_paragraph("B. [Option B]")
        doc.add_paragraph("C. [Option C]")
        doc.add_paragraph("D. [Option D]")
        doc.add_paragraph("E. [Option E]")
        doc.add_paragraph(
            f"Tags: CAH 01 > [Subtopic]; Lecture > {lecture}; Difficulty: Intermediate; AUS:3; Domain: [Curriculum area]; Yield: [High/Medium/Low]"
        )


def add_answer_key_block(doc: Document) -> None:
    doc.add_paragraph("Answer Key")
    for q in range(1, 11):
        doc.add_paragraph(f"Q{q}: [A-H]")


def add_fill_checklist(doc: Document) -> None:
    doc.add_paragraph("Pre-ingest checklist")
    doc.add_paragraph("- Replace all [placeholders].")
    doc.add_paragraph("- Keep options in A.-E. format.")
    doc.add_paragraph("- Keep question headers in Question n. format.")
    doc.add_paragraph("- Keep final Answer Key lines in Qn: X format.")


def build_template(output_path: Path, lecture: str) -> None:
    doc = Document()
    add_intro(doc, lecture)
    add_emq_block(doc, lecture)
    add_sba_block(doc, lecture)
    add_answer_key_block(doc)
    add_fill_checklist(doc)
    doc.save(output_path)


def write_readme(output_dir: Path, count: int) -> None:
    readme = output_dir / "README.md"
    lines = [
        "# Lecture DOCX Template Pack",
        "",
        "Generated parser-safe templates for CAH QBank ingestion.",
        "",
        "## How to use",
        "1. Open each lecture template and replace all placeholders.",
        "2. Keep structure unchanged (Question lines, A.-E. options, Answer Key).",
        "3. Move completed files into your content root under:",
        "   - `content/CAH_qbank/import_source/questions/`",
        "4. Run ingestion:",
        "   - `pnpm ingest`",
        "5. Verify report:",
        "   - `scripts/ingest/reports/latest.json`",
        "",
        "## Notes",
        "- You can add multiple module tags on one line separated by semicolons.",
        "- Keep MCQ-only (`SBA`, `EMQ_STEM`).",
        "",
        "## Included files",
    ]
    for i in range(1, count + 1):
        lines.append(f"- `L{lecture_code(i)}_CAH_MCQ_Template.docx`")
    readme.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate lecture DOCX templates")
    parser.add_argument("--count", type=int, default=12, help="Number of lecture templates to generate")
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("docs/templates/cah-lecture-docx-templates"),
        help="Output directory",
    )
    args = parser.parse_args()

    output_dir = args.output
    output_dir.mkdir(parents=True, exist_ok=True)

    for i in range(1, args.count + 1):
        lecture = lecture_code(i)
        file_path = output_dir / f"L{lecture}_CAH_MCQ_Template.docx"
        build_template(file_path, lecture)

    write_readme(output_dir, args.count)

    zip_base = output_dir.parent / "cah-lecture-docx-templates"
    shutil.make_archive(str(zip_base), "zip", root_dir=output_dir)

    print(f"Generated {args.count} templates in {output_dir}")
    print(f"Zip archive: {zip_base}.zip")


if __name__ == "__main__":
    main()
